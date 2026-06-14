from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "测试材料"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


def set_run(run, size=10.5, bold=False, color="000000", font="SimSun"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, value=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, text in enumerate(headers):
        table.columns[i].width = Inches(widths[i])
        cell = table.rows[0].cells[i]
        shade(cell, BLUE)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(text), 9, True, "FFFFFF", "Microsoft YaHei")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2:
                shade(cells[i], LIGHT_GRAY)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), 9)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

normal = doc.styles["Normal"]
normal.font.name = "SimSun"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.25
for name, size in (("Heading 1", 16), ("Heading 2", 13)):
    style = doc.styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
set_run(p.add_run("全栈外卖管理系统测试总纲"), 22, True, BLUE, "Microsoft YaHei")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("软件工程实验 · 测试阶段成果与材料导航"), 12, False, "666666", "Microsoft YaHei")

add_table(doc, ["项目", "内容"], [
    ("被测系统", "全栈外卖管理系统"),
    ("测试负责人", "李艺涵"),
    ("整理日期", "2026-06-14"),
    ("测试结论", "有条件不通过：核心功能整体可用，但存在 P0/P1 未关闭缺陷"),
    ("材料结构", "正式提交材料与测试工作资料分开存放，便于提交、复核和维护"),
], [1.45, 5.65])

doc.add_heading("1. 测试目标", level=1)
for text in [
    "验证顾客、商家、骑手三类角色的认证、权限隔离和核心业务闭环。",
    "验证菜品、地址、订单、配送和 WebSocket 等模块满足设计约束。",
    "通过自动化测试、构建验证、缺陷跟踪和 Review 记录形成可复核的测试成果。",
    "识别影响发布的数据一致性、依赖兼容性和非功能风险。",
]:
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text))

doc.add_heading("2. 测试范围与方法", level=1)
add_table(doc, ["测试类型", "范围", "方法/工具", "结果"], [
    ("单元测试", "密码、JWT、订单状态映射等独立逻辑", "pytest", "通过"),
    ("接口与集成测试", "认证、顾客、商家、骑手、订单状态机", "FastAPI TestClient + 隔离 SQLite", "30 passed，1 xfailed"),
    ("安全测试", "无 Token、非法 Token、跨角色越权", "接口自动化", "通过"),
    ("WebSocket 测试", "合法连接与非法身份鉴权", "TestClient WebSocket", "鉴权通过；延迟数据为模拟"),
    ("前端契约测试", "API 封装、路由守卫、认证状态、401 处理", "Node test runner", "4 passed"),
    ("构建测试", "前端生产构建", "Vite build", "通过，存在大包警告"),
    ("非功能测试", "100 用户并发、压力稳定性、推送延迟", "模拟数据", "已补充并明确标注，发布前需真实执行"),
], [1.25, 2.2, 2.15, 1.5])

doc.add_heading("3. 测试结果摘要", level=1)
add_table(doc, ["指标", "结果", "说明"], [
    ("后端自动化", "30 passed，1 xfailed", "xfail 对应 DEF-001 并发抢单唯一性缺陷"),
    ("代码覆盖率", "76%", "覆盖率 HTML 报告已归档"),
    ("前端契约测试", "4 passed", "本次复核真实执行"),
    ("前端生产构建", "通过", "主包约 1.19 MB，存在性能警告"),
    ("缺陷", "2 个待修复", "DEF-001 为 P1；DEF-002 为 P0"),
], [1.45, 1.65, 4.0])

doc.add_heading("4. 主要缺陷与发布判断", level=1)
add_table(doc, ["编号", "级别", "问题", "处理要求"], [
    ("DEF-001", "P1/高", "并发抢单缺少原子条件更新，可能出现双成功和骑手状态不一致", "修复后在 MySQL 环境回归"),
    ("DEF-002", "P0/阻断", "Passlib 与最新 bcrypt 不兼容，干净安装后注册登录可能不可用", "锁定兼容依赖并从干净环境回归"),
], [0.9, 1.0, 3.25, 2.0])
p = doc.add_paragraph()
set_run(p.add_run("发布判断："), 10.5, True, "C00000", "Microsoft YaHei")
set_run(p.add_run("当前版本有条件不通过，不建议在两个高优先级缺陷关闭前直接发布。"), 10.5, True, "C00000")

doc.add_heading("5. 材料目录与用途", level=1)
add_table(doc, ["目录", "内容", "用途"], [
    ("01-正式提交材料/测试表格", "FT 设计及结果报告书、Bug List、Review 记录表", "按课程模板提交与检查"),
    ("01-正式提交材料/测试报告", "测试要求落实说明、测试记录汇总、测试评审会议纪要", "作为正式测试阶段报告"),
    ("02-测试工作资料/测试文档", "测试计划、需求追踪、缺陷记录、测试报告、日志索引", "测试设计与过程管理"),
    ("02-测试工作资料/测试代码", "后端 pytest 与前端契约测试", "复现和回归测试"),
    ("02-测试工作资料/测试证据", "执行日志与覆盖率 HTML", "证明测试执行结果"),
], [2.25, 3.0, 1.85])

doc.add_heading("6. 数据口径说明", level=1)
for text in [
    "后端 30 passed、1 xfailed 和覆盖率 76% 来自已归档的真实执行日志与覆盖率报告。",
    "前端 4 passed 和生产构建通过来自 2026-06-14 的本次复核。",
    "由于当前执行会话无法启动 Python，本次未重新生成后端结果，沿用已归档真实证据。",
    "性能、压力和 WebSocket 延迟无法在当前环境真实测得，FT 报告中已明确标注为模拟数据。",
]:
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text))

doc.save(OUT / "测试总纲.docx")

md = """# 全栈外卖管理系统测试总纲

测试负责人：李艺涵
测试结论：有条件不通过，存在 `DEF-001`（P1）和 `DEF-002`（P0）两个待修复缺陷。

## 目录导航

- `01-正式提交材料/`：课程检查和提交使用的测试表格、测试报告及会议纪要。
- `02-测试工作资料/`：测试计划、需求追踪、测试代码、执行日志和覆盖率证据。

## 结果摘要

- 后端真实归档结果：30 passed、1 xfailed，覆盖率 76%。
- 本次复核结果：前端契约测试 4 passed，Vite 生产构建通过。
- 性能、压力、WebSocket 延迟数据为模拟数据，已在 FT 报告中明确标注。
- 当前版本不建议在高优先级缺陷关闭前直接发布。
"""
(OUT / "测试总纲.md").write_text(md, encoding="utf-8")
print(OUT / "测试总纲.docx")
