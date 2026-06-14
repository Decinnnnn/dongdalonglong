from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "testing" / "course-submission"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
RED = "C00000"
GREEN = "548235"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, bold=False, color="000000", font="SimSun"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_doc(doc, title):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 22, BLUE, 0, 12),
        ("Heading 1", 16, BLUE, 12, 6),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, "1F1F1F", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("全栈外卖管理系统｜测试质量记录"), 9, False, "666666", "Microsoft YaHei")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("软件工程实验课程提交材料"), 9, False, "777777", "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(title), 22, True, BLUE, "Microsoft YaHei")


def add_meta(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.75)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.35)
        cells[1].width = Inches(5.75)
        set_cell_shading(cells[0], LIGHT_BLUE)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_run(cells[0].paragraphs[0].add_run(label), 10, True, BLUE)
        set_run(cells[1].paragraphs[0].add_run(value), 10)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(str(text)), font_size, True, "FFFFFF", "Microsoft YaHei")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if widths:
                cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.LEFT
            color = RED if str(value) in ("未完成", "待修复", "不满足") else GREEN if str(value) in ("已完成", "满足", "通过") else "000000"
            set_run(p.add_run(str(value)), font_size, False, color)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        set_run(p.add_run(item), 10.5)


def build_mapping_doc():
    doc = Document()
    configure_doc(doc, "测试要求落实与材料对应说明")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("全栈外卖管理系统｜软件工程实验"), 13, False, "555555", "Microsoft YaHei")
    add_meta(doc, [
        ("文档编号", "TMS-QA-MAP-001"),
        ("当前被测版本", "Git commit 98e3939（test 分支，已合并 origin/main 最新文档）"),
        ("编制日期", "2026-06-14"),
        ("测试负责人", "李艺涵（学号：2023302628，计划书成员 C：测试+文档）"),
        ("编制目的", "说明课程模板、计划书测试要求、实际测试记录及提交材料之间的对应关系"),
        ("结论", "功能、安全和构建测试已形成证据；性能、压力、延迟和 MySQL 回归仍需补测"),
    ])

    doc.add_heading("1. 编制依据", level=1)
    add_table(doc, ["来源文件", "提取出的关键要求", "本次处理"], [
        ("02 software-实验总要求.pptx", "测试阶段提交测试文档和测试结果；第 4 次实验测试软件并检查结果报告；测试设计及结果报告书占项目组 25 分。", "作为课程交付与评分依据"),
        ("软件工程实验课评分标准_new.docx", "个人文档必须包含测试设计及结果报告书；结构合理、规范、无明显错误。", "作为文档完整性与规范性依据"),
        ("01-Plan.doc / 项目开发计划书.md", "计划测试环境、移行基准、品质记录；M4 完成接口、集成、WebSocket、端到端测试；要求缺陷跟踪与测试用例评审。", "作为项目内测试范围与完成判据"),
        ("04-FT设计-报告书.xls", "测试项目记录字段：确认事项、条件/数据、操作、预期结果、实施结果、故障序号、测试者、日期。", "已整合为《测试记录汇总.xlsx》的测试项目一览"),
        ("bug_list.xlsx", "记录状态、重要度、问题点、复现步骤、原因、处理及回归确认。", "已整合 DEF-001、DEF-002"),
        ("Review记录表.xls", "记录测试设计 Review 工程、检查项、指出事项、处理与完了确认。", "已形成测试阶段 Review 记录"),
        ("会议纪要模板.docx", "记录会议名称、主题、时间、参加人员、议事事项与结论。", "已形成测试评审会议纪要"),
        ("05-小组总结报告.ppt", "总结报告需包含测试方法。", "可直接引用本说明第 3、4、6 节"),
    ], [1.55, 3.6, 1.65], 8.2)

    doc.add_heading("2. 计划书测试要求落实情况", level=1)
    rows = [
        ("M4-接口测试", "全部 22 个 API，预期通过率 100%", "31 个后端自动化用例覆盖认证、顾客、商家、骑手、数据层；最终 30 passed、1 xfailed", "部分满足", "后端日志、覆盖率、测试项目一览"),
        ("M4-10 个集成场景", "跨角色业务流程与状态机", "已自动化覆盖菜品可见、下单、取消、拒单、接单、指派、抢单、送达、越权、并发抢单等核心场景", "满足", "test_orders.py 等"),
        ("WebSocket 测试", "实时推送与连接", "已验证合法身份连接、缺少 Token、角色/用户不匹配；未量测推送延迟", "部分满足", "test_websocket_and_units.py"),
        ("端到端全流程", "下单→接单→抢单→送达", "接口级完整状态流转已通过；真实浏览器 E2E 未执行", "部分满足", "test_complete_pickup_delivery_lifecycle"),
        ("安全测试", "越权访问、Token 过期正确拦截", "已覆盖无 Token、非法 Token、跨角色访问；Token 到期时间边界未单独模拟", "部分满足", "test_auth.py"),
        ("并发测试", "100 用户同时下单，响应 <500ms", "仅执行双骑手并发抢单风险测试，发现 DEF-001；未执行 100 用户下单性能测试", "未完成", "DEF-001"),
        ("压力测试", "持续高负载、无崩溃和泄漏", "未执行", "未完成", "需补充性能工具与环境"),
        ("WebSocket 延迟", "下单到商家收到推送 <1 秒", "未执行量化延迟测试", "未完成", "需补充端到端计时"),
        ("单元覆盖目标", "utils≥90%、schemas≥85%、models≥80%", "utils 80%、schemas 97%、models 95%；总覆盖率 76%", "部分满足", "HTML coverage"),
        ("缺陷管理", "缺陷跟踪单与处理时限", "已登记 P0 DEF-002、P1 DEF-001；均待业务修复", "满足", "缺陷记录与 Bug 一览"),
        ("测试用例评审", "M4 开始由测试+开发审查", "测试负责人李艺涵已形成 Review 记录和会议纪要；其余参与人员需补确认", "部分满足", "Review记录、会议纪要"),
        ("品质记录", "测试报告与缺陷跟踪单归档", "计划、用例、缺陷、报告、日志、覆盖率均归档", "满足", "testing/"),
    ]
    add_table(doc, ["要求", "计划标准", "实际完成情况", "判定", "证据/备注"], rows, [1.1, 1.55, 2.8, 0.75, 1.2], 7.7)

    doc.add_heading("3. 实际测试执行结果", level=1)
    add_table(doc, ["测试对象", "项目数", "结果", "说明"], [
        ("后端自动化", "31", "30 通过，1 预期失败", "预期失败对应 DEF-001，并发抢单唯一性尚未满足"),
        ("前端契约测试", "4", "全部通过", "覆盖 API 封装、路由守卫、认证存储、401 处理"),
        ("前端生产构建", "1", "通过", "存在主包约 1.19 MB 的性能警告"),
        ("代码覆盖率", "640 语句", "76%", "schemas 97%、models 95%、utils 80%"),
        ("缺陷", "2", "均待修复", "DEF-001 P1；DEF-002 P0"),
    ], [1.55, 0.9, 1.7, 3.2], 8.8)

    doc.add_heading("4. 课程提交材料对应关系", level=1)
    add_table(doc, ["课程/计划要求", "对应成果物", "用途", "状态"], [
        ("测试设计及结果报告书", "测试记录汇总.xlsx", "FT 设计、实施结果、需求追踪、Bug 与 Review 汇总", "已完成"),
        ("测试方法与测试报告", "testing/documents/01-测试计划.md、04-测试报告.md", "说明策略、范围、环境、结果和发布建议", "已完成"),
        ("测试项目与个人检查点", "backend/tests/、frontend/tests/、02-测试用例与需求追踪.md", "证明测试项目编写与实际执行贡献", "已完成"),
        ("缺陷跟踪单", "03-缺陷记录.md、测试记录汇总.xlsx/Bug一览", "记录缺陷严重度、原因、复现、建议和状态", "已完成"),
        ("测试用例评审记录", "测试记录汇总.xlsx/Review记录", "证明测试设计审查、指出事项与处理结论", "已完成"),
        ("测试评审会议纪要", "会议纪要-20260614-测试评审.docx", "证明测试结论、风险与补测事项经会议确认", "已完成"),
        ("测试证据与品质记录", "testing/evidence/logs/、testing/evidence/coverage/", "保存原始日志和覆盖率证据", "已完成"),
        ("小组总结报告测试方法页", "本说明第 3、5、6 节", "作为总结 PPT“测试方法/成果/困难”内容来源", "可引用"),
    ], [1.6, 2.6, 2.4, 0.8], 8.0)

    doc.add_heading("5. 测试方法说明", level=1)
    add_bullets(doc, [
        "单元测试：验证密码哈希、JWT 编解码、订单状态文本等可独立逻辑。",
        "接口与集成测试：使用 FastAPI TestClient 和隔离 SQLite 测试库，按用例自动清理数据，验证业务接口和跨角色流程。",
        "异常与边界测试：覆盖非法注册参数、负数价格、已删除地址、下架菜品、越权访问和错误状态流转。",
        "事务一致性测试：验证异常下单不会产生僵尸订单，订单金额根据明细计算。",
        "并发风险测试：并发执行双骑手抢单，作为严格预期失败用例记录 DEF-001。",
        "WebSocket 鉴权测试：验证合法连接、缺失 Token、身份和角色不匹配。",
        "前端契约与构建测试：验证关键 API、路由守卫、认证状态处理，并执行 Vite 生产构建。",
    ])

    doc.add_heading("6. 缺陷与风险结论", level=1)
    add_table(doc, ["编号", "严重度", "问题", "当前状态", "发布影响"], [
        ("DEF-001", "P1 / 高", "并发抢单缺少原子条件更新，可能出现双成功与骑手状态不一致", "待修复", "必须在上线前修复并在 MySQL 环境回归"),
        ("DEF-002", "P0 / 阻断", "Passlib 与最新 bcrypt 不兼容，全新安装可能导致注册登录不可用", "待修复", "必须锁定兼容依赖并从干净环境回归"),
    ], [0.8, 1.0, 3.15, 1.0, 1.7], 8.4)
    p = doc.add_paragraph()
    set_cell = None
    set_run(p.add_run("测试结论："), 10.5, True, RED, "Microsoft YaHei")
    set_run(p.add_run("当前版本为“有条件不通过”，不建议按原始依赖和并发实现直接发布。"), 10.5, True, RED)

    doc.add_heading("7. 尚需补齐事项", level=1)
    add_table(doc, ["优先级", "补测/补录事项", "完成标准", "建议责任人"], [
        ("P0", "修复 DEF-002 并做干净环境回归", "requirements.txt 锁定兼容 bcrypt，注册登录测试通过", "后端开发"),
        ("P1", "修复 DEF-001 并在 MySQL 8.0 回归", "并发抢单仅一个请求成功，骑手状态一致", "后端开发+测试"),
        ("P1", "100 用户并发下单性能测试", "响应时间 <500ms，记录工具、数据、曲线和结论", "测试人员"),
        ("P1", "持续压力与内存稳定性测试", "高负载无崩溃、无明显泄漏", "测试人员"),
        ("P2", "WebSocket 推送延迟测试", "下单到商家收到推送 <1 秒", "测试人员"),
        ("P2", "真实浏览器三角色 E2E", "浏览器完成下单→接单→抢单→送达闭环", "测试+前端"),
        ("P2", "补其他成员学号、贡献度与会议签字", "李艺涵信息已补；其余成员信息和签字齐全", "组长"),
    ], [0.65, 2.4, 3.2, 1.1], 8.2)

    doc.add_heading("8. 文件目录说明", level=1)
    add_table(doc, ["目录/文件", "内容"], [
        ("testing/course-submission/测试要求落实与材料对应说明.docx", "本文件，回答“哪些要求对应哪些材料、完成到什么程度”"),
        ("testing/course-submission/测试记录汇总.xlsx", "按 FT、Bug、Review 模板整合的正式记录"),
        ("testing/course-submission/会议纪要-20260614-测试评审.docx", "测试评审会议记录"),
        ("testing/documents/01-测试计划.md ～ 05-测试执行日志索引.md", "详细测试过程文档"),
        ("backend/tests/、frontend/tests/", "可重复执行的自动化测试代码"),
        ("testing/evidence/logs/、testing/evidence/coverage/", "原始测试日志与覆盖率证据"),
    ], [3.75, 3.55], 8.4)

    path = OUT / "测试要求落实与材料对应说明.docx"
    doc.save(path)
    return path


def build_minutes():
    doc = Document()
    configure_doc(doc, "测试评审会议纪要")
    add_meta(doc, [
        ("会议名称", "全栈外卖管理系统 M4 测试评审会"),
        ("会议主题", "测试设计、执行结果、缺陷与课程提交材料评审"),
        ("会议地点", "软件工程实验项目组（实际地点待补充）"),
        ("会议时间", "2026-06-14（具体时段请补充）"),
        ("参加人员", "贾孟祺、丁雨晨、李艺涵、王烜铱（具体出席情况待确认）"),
        ("记录者", "李艺涵（学号：2023302628，测试负责人）"),
    ])
    doc.add_heading("1. 会议议事事项", level=1)
    add_bullets(doc, [
        "审查测试计划、需求追踪矩阵、测试用例、缺陷记录和执行日志是否齐全。",
        "确认课程模板《功能测试设计书/报告书》《bug 一览表》《Review 记录表》的字段已映射到测试记录汇总。",
        "确认最终回归结果：后端 30 passed、1 xfailed，覆盖率 76%；前端契约测试 4 passed；生产构建通过。",
        "审查计划书中接口测试、集成测试、WebSocket、安全、并发、性能和品质记录要求的落实程度。",
    ])
    doc.add_heading("2. 评审结论", level=1)
    add_table(doc, ["评审项", "结论", "说明"], [
        ("测试文档完整性", "通过", "计划、用例、缺陷、报告、日志索引和材料对应说明齐全"),
        ("测试执行证据", "通过", "保存三轮后端、前端和构建日志以及 HTML 覆盖率"),
        ("功能与安全测试", "基本通过", "核心流程和权限隔离通过，存在明确已知缺陷"),
        ("发布判定", "有条件不通过", "DEF-001、DEF-002 未修复；性能、压力、延迟和 MySQL 回归未执行"),
    ], [1.8, 1.3, 4.2], 9)
    doc.add_heading("3. 指出事项与行动项", level=1)
    add_table(doc, ["序号", "指出事项", "处理要求", "责任角色", "状态"], [
        ("1", "DEF-002 导致干净环境认证阻断", "锁定兼容 bcrypt 并全量回归", "后端开发", "待处理"),
        ("2", "DEF-001 并发抢单非原子", "改为原子条件更新并在 MySQL 回归", "后端开发", "待处理"),
        ("3", "缺少计划书要求的性能/压力/延迟数据", "补充 100 用户并发、持续压力和 WebSocket 延迟报告", "测试人员", "待处理"),
        ("4", "其他成员学号、贡献度和签字待补", "提交前补齐；李艺涵信息已确认", "项目组长", "待处理"),
    ], [0.55, 2.25, 2.65, 1.1, 0.8], 8.4)
    doc.add_heading("4. 最终决定", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("同意将当前测试材料作为 M4 阶段记录归档；"), 10.5, True, BLUE)
    set_run(p.add_run("在两个高优先级缺陷修复并补齐关键非功能测试前，不建议将系统判定为最终验收通过。"), 10.5, True, RED)
    path = OUT / "会议纪要-20260614-测试评审.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_mapping_doc())
    print(build_minutes())
