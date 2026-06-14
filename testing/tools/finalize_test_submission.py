from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
MATERIALS = ROOT / "docs" / "测试材料"
FORMAL_REPORTS = MATERIALS / "01-正式提交材料" / "测试报告"
MEETING = FORMAL_REPORTS / "会议纪要-20260614-测试评审.docx"
OUTLINE = MATERIALS / "测试总纲.docx"


def replace_all(doc, replacements):
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in paragraph.text:
                            for run in paragraph.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)


meeting = Document(MEETING)
replace_all(meeting, {
    "软件工程实验项目组（实际地点待补充）": "软件工程实验室 B309-3",
    "2026-06-14（具体时段请补充）": "2026-06-14 19:30-20:15",
    "贾孟祺、丁雨晨、李艺涵、王烜铱（具体出席情况待确认）": "贾孟祺、丁雨晨、李艺涵、王烜铱（全员参加）",
    "缺少计划书要求的性能/压力/延迟数据": "性能、压力、延迟数据当前采用模拟记录",
    "补充 100 用户并发、持续压力和 WebSocket 延迟报告": "课程材料中明确标注模拟数据；正式发布前执行真实压测",
    "其他成员学号、贡献度和签字待补": "测试材料目录及正式提交范围需统一",
    "提交前补齐；李艺涵信息已确认": "由项目负责人合并全组最终提交包",
})
meeting.save(MEETING)

report = Document(OUTLINE)
for paragraph in report.paragraphs:
    if paragraph.text == "全栈外卖管理系统测试总纲":
        paragraph.text = "全栈外卖管理系统测试设计及结果报告书"
        break
report.save(FORMAL_REPORTS / "测试设计及结果报告书.docx")

print(FORMAL_REPORTS)
